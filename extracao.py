import win32com.client
import pythoncom
import os
import fitz  # PyMuPDF
from docx import Document
import re
from dotenv import load_dotenv

load_dotenv()  # Carrega variáveis do .env


def baixar_pdfs_outlook(remetente, conta_email):
    pythoncom.CoInitialize()
    try:
        outlook = win32com.client.Dispatch("Outlook.Application").GetNamespace("MAPI")
        pasta_destino = os.path.dirname(os.path.abspath(__file__))
        arquivos_salvos = []
        for conta in outlook.Folders:
            if conta.Name.lower() == conta_email.lower():
                inbox = conta.Folders["Caixa de Entrada"]
                emails = inbox.Items.Restrict(f"[SenderEmailAddress] = '{remetente}'")
                emails.Sort("[ReceivedTime]", True)
                for item in emails:
                    for anexo in item.Attachments:
                        if anexo.FileName.endswith(".pdf"):
                            nome_arquivo = os.path.join(pasta_destino, anexo.FileName)
                            anexo.SaveAsFile(nome_arquivo)
                            arquivos_salvos.append(nome_arquivo)
                            print(f"PDF salvo de {remetente}: {nome_arquivo}")
        return arquivos_salvos
    finally:
        pythoncom.CoUninitialize()


def extrair_texto_pdf(caminho_pdf):
    texto = ""
    with fitz.open(caminho_pdf) as pdf:
        for pagina in pdf:
            texto += pagina.get_text()
    return texto


def buscar_pdfs_pasta(pasta):
    return [os.path.join(pasta, arquivo) for arquivo in os.listdir(pasta) if arquivo.endswith(".pdf")]


def comparar_pdfs(lista_pdfs_email, pasta_pdfs):
    lista_pdfs_pasta = buscar_pdfs_pasta(pasta_pdfs)
    if len(lista_pdfs_email) != len(lista_pdfs_pasta):
        print("⚠ Número de PDFs no e-mail e na pasta não corresponde. Verifique os arquivos.")
    resultados = []
    for pdf_email, pdf_pasta in zip(lista_pdfs_email, lista_pdfs_pasta):
        autenticidade = "Autêntico" if extrair_texto_pdf(pdf_email) == extrair_texto_pdf(pdf_pasta) else "Falso"
        resultados.append((pdf_email, pdf_pasta, autenticidade))
    return resultados


def gerar_oficio(autenticidade, pdf_recebido, juiz_destino, processo_numero):
    pasta_destino = os.path.dirname(os.path.abspath(__file__))
    caminho_oficio = os.path.join(pasta_destino, "Oficio_Resposta.docx")
    doc = Document()
    doc.add_paragraph("PODER JUDICIÁRIO")
    doc.add_paragraph("Tribunal de Justiça")
    doc.add_paragraph("\n")
    doc.add_paragraph(f"Ao Juízo responsável: {juiz_destino}")
    doc.add_paragraph(f"Processo nº {processo_numero}")
    doc.add_paragraph("\n")
    doc.add_paragraph("Assunto: Autenticidade de documento recebido")
    doc.add_paragraph("\n")
    doc.add_paragraph(f"Informamos que o documento '{pdf_recebido}' enviado para verificação foi analisado.")
    doc.add_paragraph(f"Resultado da análise: **{autenticidade}**.")
    doc.add_paragraph("\n")
    doc.add_paragraph("Colocamo-nos à disposição para quaisquer esclarecimentos adicionais.")
    doc.add_paragraph("\n")
    doc.add_paragraph("Atenciosamente,")
    doc.add_paragraph("Setor de Verificação Documental")
    doc.add_paragraph("Data: _________")
    doc.save(caminho_oficio)
    print(f"Ofício gerado com sucesso: {caminho_oficio}")
    return caminho_oficio


def extrair_informacoes_pdf(caminho_pdf):
    texto = extrair_texto_pdf(caminho_pdf)
    juiz_pattern = re.search(r"Juiz(?:[\s:]+)(.+)", texto, re.IGNORECASE)
    processo_pattern = re.search(r"Processo nº\s*([\d\-\.]+)", texto)
    juiz = juiz_pattern.group(1).strip() if juiz_pattern else "Não encontrado"
    processo = processo_pattern.group(1).strip() if processo_pattern else "Não encontrado"
    return juiz, processo


# --- Fluxo principal ---
remetente = os.environ.get("REMETENTE_PADRAO", "")
conta_email = os.environ.get("CONTA_EMAIL", "")
pasta_pdf = "pdfs"

pdfs_recebidos = baixar_pdfs_outlook(remetente, conta_email)
if pdfs_recebidos:
    resultado_comparacao = comparar_pdfs(pdfs_recebidos, pasta_pdf)
    for pdf_email, pdf_pasta, autenticidade in resultado_comparacao:
        print(f"{pdf_email} vs {pdf_pasta} → {autenticidade}")
    juiz, processo = extrair_informacoes_pdf(pdfs_recebidos[0])
    caminho_oficio = gerar_oficio(resultado_comparacao[0][2], pdfs_recebidos[0], juiz, processo)
    print(f"Ofício gerado em: {caminho_oficio}")
else:
    print("Nenhum PDF recebido para análise ou geração de ofício.")

